from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "Ezviz-Office项目下载与安装教程.docx"


def set_cn_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F5F7")
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run(text)
    set_cn_font(r, "Consolas", 9, color=(32, 45, 55))


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cn_font(r)


doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(1.8), Cm(1.8)
sec.left_margin, sec.right_margin = Cm(2.1), Cm(2.1)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Title", 24, (31, 78, 121)), ("Heading 1", 16, (31, 78, 121)), ("Heading 2", 12, (55, 95, 145))]:
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(*color)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Ezviz Office 项目下载与安装教程")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Windows 环境 · GitHub 下载 · 依赖安装 · 启动与更新")
set_cn_font(r, size=11, color=(100, 100, 100))

doc.add_heading("一、安装基础环境", level=1)
doc.add_heading("1. 安装 Git", level=2)
p = doc.add_paragraph("从 Git 官网下载安装：")
add_hyperlink(p, "https://git-scm.com/download/win", "https://git-scm.com/download/win")
doc.add_paragraph("安装完成后打开 PowerShell，检查是否成功：")
code(doc, "git --version")

doc.add_heading("2. 安装 Node.js", level=2)
p = doc.add_paragraph("建议安装 Node.js LTS 版本：")
add_hyperlink(p, "https://nodejs.org/", "https://nodejs.org/")
doc.add_paragraph("安装完成后检查：")
code(doc, "node --version\nnpm --version")

doc.add_heading("二、从 GitHub 下载项目", level=1)
doc.add_heading("方法一：使用 Git 下载（推荐）", level=2)
doc.add_paragraph("在准备存放项目的文件夹中打开 PowerShell，依次运行：")
code(doc, "git clone https://github.com/zhangdashuais/Ezviz-office.git\ncd Ezviz-office")
doc.add_paragraph("以后同步 GitHub 最新版本：")
code(doc, "git pull origin main")
doc.add_paragraph("如需强制使用 GitHub 版本覆盖本地已跟踪文件：")
code(doc, "git fetch origin\ngit reset --hard origin/main")
p = doc.add_paragraph()
r = p.add_run("注意：reset --hard 会丢弃本地未提交的修改。")
set_cn_font(r, bold=True, color=(192, 57, 43))

doc.add_heading("方法二：下载 ZIP", level=2)
for item in ["打开项目 GitHub 页面。", "点击绿色的 Code 按钮。", "选择 Download ZIP。", "下载完成后解压，并在解压目录中打开 PowerShell。"]:
    bullet(doc, item)
p = doc.add_paragraph("项目地址：")
add_hyperlink(p, "https://github.com/zhangdashuais/Ezviz-office", "https://github.com/zhangdashuais/Ezviz-office")
doc.add_paragraph("ZIP 方式不能直接使用 git pull 更新，后续需重新下载或改用 Git。")

doc.add_heading("三、安装项目依赖", level=1)
doc.add_paragraph("进入项目根目录后运行：")
code(doc, "npm install")
doc.add_paragraph("安装 Playwright 使用的 Chromium 浏览器：")
code(doc, "npx playwright install chromium")
doc.add_paragraph("项目主要依赖：")
for item in ["Express：本地服务", "Multer：文件上传", "Playwright：浏览器自动化"]:
    bullet(doc, item)

doc.add_heading("四、启动项目", level=1)
code(doc, "npm start")
doc.add_paragraph("启动成功后，在浏览器打开：")
p = doc.add_paragraph()
add_hyperlink(p, "http://localhost:3217/inline-packager.html", "http://localhost:3217/inline-packager.html")
doc.add_paragraph("停止项目时，在运行窗口按 Ctrl + C。")

doc.add_heading("五、账号密码文件", level=1)
doc.add_paragraph("商城后台自动配置需要网站账号密码 Excel，该文件不会存入 GitHub。程序会优先在当前 Windows 用户桌面查找文件名中包含以下文字的 .xlsx 文件：")
bullet(doc, "网站账号密码")
bullet(doc, "账号密码")
doc.add_paragraph("也可以通过环境变量指定文件夹：")
code(doc, '$env:EZVIZ_CREDENTIAL_DIR="D:\\账号文件目录"\nnpm start')

doc.add_heading("六、外部巡查功能依赖", level=1)
doc.add_paragraph("Banner / Popup 的部分巡查功能还依赖以下外部目录：")
code(doc, "E:\\Website-backend\\backend-operations")
doc.add_paragraph("主要需要：")
code(doc, "E:\\Website-backend\\backend-operations\\website-audit\\config\\banner-check.json\nE:\\Website-backend\\backend-operations\\website-audit\\scripts\\check-homepage-campaign-rendered.mjs")
doc.add_paragraph("如果只使用普通页面、WTB 或部分产品处理功能，可以先启动项目；使用 Banner / Popup 巡查时需准备上述目录。")

doc.add_heading("七、常见问题", level=1)
doc.add_heading("npm 无法识别", level=2)
doc.add_paragraph("重新打开 PowerShell。如果仍无法识别，请重新安装 Node.js，并确认安装程序已将 Node.js 加入 PATH。")
doc.add_heading("端口 3217 被占用", level=2)
code(doc, '$env:PORT="3218"\nnpm start')
doc.add_paragraph("随后访问 http://localhost:3218/inline-packager.html。")
doc.add_heading("Playwright 找不到浏览器", level=2)
code(doc, "npx playwright install chromium")
doc.add_heading("更新后依赖发生变化", level=2)
code(doc, "npm install\nnpx playwright install chromium\nnpm start")

doc.add_heading("八、首次安装完整命令", level=1)
code(doc, "git clone https://github.com/zhangdashuais/Ezviz-office.git\ncd Ezviz-office\nnpm install\nnpx playwright install chromium\nnpm start")
doc.add_paragraph("启动后访问：http://localhost:3217/inline-packager.html")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Ezviz Office · 项目安装指南")
set_cn_font(r, size=8, color=(120, 120, 120))

doc.save(OUT)
print(OUT)
