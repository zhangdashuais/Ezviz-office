from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from pathlib import Path
from datetime import date

OUT = Path(__file__).resolve().parents[1] / "Ezviz-Office功能介绍说明.docx"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F5F7"


def font(run, name="Microsoft YaHei", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell_or_para, fill):
    target = cell_or_para._tc.get_or_add_tcPr() if hasattr(cell_or_para, "_tc") else cell_or_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    target.append(shd)


def add_text(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text), bold=bold, color=color)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(item))


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(item))


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    shade(p, LIGHT_GRAY)
    font(p.add_run(text), "Consolas", 9, color="20303D")


def info_box(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.margin_top = cell.margin_bottom = 100
    p = cell.paragraphs[0]
    font(p.add_run(title + "："), bold=True, color=BLUE)
    font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Cm(4.2), Cm(6.3), Cm(5.5)]
    for i, title in enumerate(["功能模块", "主要用途", "执行性质"]):
        cell = table.rows[0].cells[i]
        shade(cell, BLUE)
        cell.width = widths[i]
        font(cell.paragraphs[0].add_run(title), bold=True, color="FFFFFF")
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].width = widths[i]
            font(cells[i].paragraphs[0].add_run(value), size=9.5)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(1.8)
sec.left_margin = sec.right_margin = Cm(2.0)

for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
doc.styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Title", 25, BLUE), ("Heading 1", 17, BLUE), ("Heading 2", 13, "375F91"), ("Heading 3", 11, "4F81BD")]:
    doc.styles[name].font.size = Pt(size)
    doc.styles[name].font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(title.add_run("Ezviz Office 功能介绍说明"), size=25, bold=True, color=BLUE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(sub.add_run("本地网站运营与办公自动化平台"), size=12, color="666666")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(meta.add_run(f"文档版本：1.0　更新日期：{date.today().isoformat()}"), size=9, color="888888")

doc.add_heading("1. 项目概述", level=1)
add_text(doc, "Ezviz Office 是一个运行在 Windows 本机的办公自动化平台，面向 EZVIZ 网站内容制作、资料处理、商城后台配置与站点质量巡查。用户通过统一网页界面上传文件、填写参数或选择站点，系统在本地完成内容转换，并在需要时调用真实浏览器进入后台执行操作。")
info_box(doc, "使用原则", "带有“生成、预览、清单”字样的操作通常不会提交后台；带有“执行、上传、发布、配置”字样的操作可能修改真实后台数据，执行前应核对站点、文件和链接。")

doc.add_heading("2. 启动与访问", level=1)
add_text(doc, "在项目目录中打开 PowerShell：")
code(doc, "npm start")
add_text(doc, "启动后访问：")
code(doc, "http://localhost:3217/inline-packager.html")
bullets(doc, ["后台自动化可能打开真实浏览器，执行期间不要关闭该窗口。", "账号密码可从网站账号密码 Excel 自动读取，也可在页面手动覆盖。", "页面下方的状态和结果区域用于查看进度、成功结果与错误原因。"])

doc.add_heading("3. 功能总览", level=1)
feature_table(doc, [
    ("一键内联打包", "把网页 HTML、CSS 和图片素材整理为可上线的内联 HTML", "本地生成"),
    ("PDF 上传转地址", "上传 PDF 并返回可复制的线上文件地址", "上传文件"),
    ("Spec 参数解析", "把规格 Excel 转成 PC / Mobile Specifications HTML", "本地生成"),
    ("文字翻译 / i18n", "提取页面文案并生成或复用 i18n key", "本地生成 / 导出"),
    ("服务中心资料上传", "创建下载资料、扩展语言、更新产品图并生成归档计划", "可提交后台"),
    ("语言包上传", "上传 XLS/XLSX 语言包到商城 Language Management", "提交后台"),
    ("WTB 产品购买链接", "按产品和渠道配置 Where To Buy 购买链接", "提交后台"),
    ("Banner / Popup", "创建活动资源、上传图片、发布并校验 UTM", "提交后台"),
    ("活动巡查", "检查 Banner / Popup 坏链和 UTM 命名", "读取 / 可修复 UTM"),
    ("EZVIZ 官网巡查", "随机抽取多站点产品并检查内容质量", "读取巡查"),
    ("Datasheet 生成 Spec 表", "从 PDF 或规格文本生成统一 Spec Excel", "本地生成"),
    ("Specification 翻译上架", "翻译并保存产品 Specifications 内容", "可提交后台"),
    ("后台产品替换", "只读 Details 下的 Overview 与 Specifications HTML", "已开放只读接口"),
])

doc.add_heading("4. 内容制作与文件处理", level=1)
doc.add_heading("4.1 一键内联打包", level=2)
add_text(doc, "用于将活动页或产品页项目中的 HTML、CSS、图片引用整理成单个可上线 HTML。")
bullets(doc, ["选择项目文件夹及目标 HTML 文件。", "设置 CSS 兼容模式和图片处理模式。", "可指定图片基础地址和文件上传接口。", "点击“开始一键内联”，在输出区复制处理后的 HTML。"])

doc.add_heading("4.2 PDF 上传转地址", level=2)
add_text(doc, "支持一次选择一个或多个 PDF，上传后生成线上地址列表。")
steps(doc, ["选择 PDF 文件。", "确认上传接口。", "点击“开始上传”。", "上传成功后点击“复制地址”。"])

doc.add_heading("4.3 Spec 参数解析", level=2)
add_text(doc, "读取规格 Excel，并根据语言、标题、图片等配置生成 PC 和 Mobile 两套 Specifications HTML。")
bullets(doc, ["输入：Excel、固定参数数量、语言、标题、图片 URL 和 Alt。", "输出：PC HTML 与 Mobile HTML。", "可分别复制两套输出，用于不同端页面。"])

doc.add_heading("4.4 Datasheet 生成 Spec 表", level=2)
add_text(doc, "从 datasheet PDF 或手动粘贴的 Specifications 文本中提取规格，生成统一结构的 Spec Excel。")
steps(doc, ["上传 PDF 或粘贴 Specifications 文本。", "选择产品类型并填写型号。", "设置输出文件名。", "预览解析结果并生成、下载 Excel。"])

doc.add_heading("5. 翻译与语言管理", level=1)
doc.add_heading("5.1 文字翻译 / i18n", level=2)
add_text(doc, "用于从 HTML 或文本中提取可翻译文案，生成 i18n key，并支持复用已上传语言表中 value 完全一致的已有 key。")
bullets(doc, ["设置语言 Key 前缀，例如 goods.Alarm_light。", "解析并提取文案。", "已有语言表中相同文案优先复用原 key。", "同一轮重复的新文案共用同一个新 key。", "完成后导出 i18n 文件。"])

doc.add_heading("5.2 语言包上传", level=2)
add_text(doc, "把 .xls 或 .xlsx 语言包上传至商城后台 Language Management。")
steps(doc, ["加载站点并只选择一个目标站点。", "选择语言包文件。", "填写语言代码；留空时尝试从文件名开头识别。", "点击“上传语言包”，查看后台响应和日志。"])

doc.add_heading("5.3 Specification 翻译上架", level=2)
add_text(doc, "用于生成产品 Specification HTML，并在确认站点、产品和语言后保存至相应后台。执行前应先检查生成内容和产品定位是否正确。")

doc.add_heading("6. 服务中心资料上传", level=1)
add_text(doc, "用于处理产品 Datasheet、Spec、高清图等资料，并按勾选动作创建下载中心资料、访问语言扩展接口、更新产品背景图或生成 SharePoint 归档计划。")
steps(doc, ["填写产品标题和关联产品搜索词。", "选择文件类型、状态和权重。", "选择产品资料文件夹或零散文件。", "核对系统识别出的 Datasheet、Spec 和高清图。", "勾选需要执行的后台动作。", "点击“开始执行”，在日志中核对每一步结果。"])
info_box(doc, "文件命名建议", "文件名包含 datasheet、spec、高清图等关键词，有助于系统正确识别资料类型。SharePoint 功能当前以生成归档计划为主。")

doc.add_heading("7. WTB 产品购买链接", level=1)
add_text(doc, "WTB 模块用于配置产品编辑页 Additional Information 中的 wheretobuy Settings。系统按 Product 定位产品，再按 Channel 找到渠道字段并写入 Purchasing Link。")
doc.add_heading("7.1 Excel 模板", level=2)
feature_table(doc, [
    ("Product", "后台产品名称，例如 RS20 Pro", "必填"),
    ("Product Page URL", "产品前台页面，用于自动识别站点和保存后复查", "建议填写"),
    ("Channel", "购买渠道，例如 Shopee、Lazada", "必填"),
    ("Purchasing Link", "对应渠道的购买链接", "必填"),
])
add_text(doc, "程序读取 Excel 中全部工作表，只处理已填写 Channel 或 Purchasing Link 的行。")
doc.add_heading("7.2 站点识别规则", level=2)
bullets(doc, ["优先解析 Product Page URL 的站点路径，例如 /id/ 自动识别为 Indonesia。", "URL 无法识别时，使用页面手动勾选的单个站点。", "如果同一文件包含多个站点，系统停止并要求按站点拆分。"])
doc.add_heading("7.3 执行流程", level=2)
steps(doc, ["上传 Excel 或填写单条测试数据。", "点击“生成 WTB 清单”检查产品、站点、渠道和链接，不提交后台。", "点击“执行 WTB 后台配置”。", "系统登录商城、搜索产品并打开编辑页。", "进入 Additional Information，匹配渠道链接字段并填写。", "点击 Complete，由后台页面发送保存请求。", "打开 Product Page URL，检查前台渠道或链接是否出现。"])
info_box(doc, "安全提示", "测试链接也会保存到真实后台。提交前应检查 Purchasing Link；如仅验证识别结果，请使用“生成 WTB 清单”。", "FCE4D6")

doc.add_heading("8. Banner / Popup 与活动巡查", level=1)
doc.add_heading("8.1 Banner", level=2)
add_text(doc, "用于上传 PC / Mobile Banner 图片、填写标题与跳转链接、调整轮播位置、保存并按需发布。普通站点默认插入第 1 位，波兰站默认插入第 2 位。")
bullets(doc, ["支持标题、Slogan、Model、Introduction、字体颜色、上下线时间。", "可隐藏 More 按钮、设置新窗口打开并选择上传后发布。", "内部 EZVIZ 链接会按站点和实际位置修正 UTM。", "发布后自动巡查；UTM 问题可自动修复，坏链只报告。"])
code(doc, "utm_source={siteCode}_banner\nutm_medium=banner{position}\nutm_campaign=web_{siteCode}_banner")

doc.add_heading("8.2 Popup", level=2)
add_text(doc, "用于创建活动 Popup，配置展示位置、频率、时间、Web/Mobile 跳转地址和图片。")
bullets(doc, ["内部 EZVIZ 链接自动应用 Popup UTM，外部链接保持原样。", "当前按单资源位处理：旧 Popup 已过期时先删除再创建；未过期时停止并提示。", "可选择提交后立即启用。"])
code(doc, "utm_source={siteCode}_popup\nutm_medium=popup\nutm_campaign=web_{siteCode}_popup")

doc.add_heading("8.3 Banner / Popup 巡查", level=2)
add_text(doc, "按所选站点检查首页活动链接的可达性和 UTM 命名。")
bullets(doc, ["巡查类型支持 banner、popup、banner + popup。", "结果包含站点、资源类型、位置、状态码、当前链接及建议修正 URL。", "坏链仅报告，不自动替换。", "Banner UTM 可通过修复功能写回后台。"])

doc.add_heading("9. EZVIZ 官网随机产品巡查", level=1)
add_text(doc, "面向多个 EZVIZ 区域站点执行产品内容质量检查。系统从 Security Cameras 与 Smart Home 分类随机抽取产品，并持续返回任务进度和问题清单。")
bullets(doc, ["检查产品标语语言是否与站点一致。", "检查 Detail、Specifications、Support 导航是否完整。", "检查 Detail 与 Specifications 内容语言。", "检查 Support 页面可访问性、Datasheet 和图片加载。", "结果按站点、产品和问题类型汇总。"])

doc.add_heading("10. 操作安全与故障处理", level=1)
bullets(doc, ["真实提交前先使用“生成清单”或“预览”检查数据。", "一次后台配置只处理一个目标站点，避免账号与站点不匹配。", "自动化浏览器运行时不要手动关闭或频繁切换页面。", "如提示登录失败，检查账号密码 Excel、站点域名及当前登录账号。", "如 Playwright 浏览器缺失，运行 npx playwright install chromium。", "如端口 3217 被占用，可设置 PORT 环境变量后重新启动。", "输出区日志应保留到问题定位完成，其中包含产品定位、保存状态和复查结果。"])

doc.add_heading("11. 当前限制", level=1)
bullets(doc, ["后台产品替换当前只开放 Overview / Specifications 读取，不执行写入或替换保存。", "Banner / Popup 巡查配置和脚本位于项目内，运行时仍需要全局可用的 Playwright CLI。", "SharePoint 资料处理目前以生成归档计划为主。", "后台页面结构或接口变化可能导致自动化控件无法匹配，需要根据日志更新规则。"])

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Ezviz Office · 功能介绍说明"), size=8, color="888888")

doc.save(OUT)
print(OUT)
